# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Freight_Intelligence_SSO_Server_Deployment_Guide.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(100, 116, 139)
HEADER_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
WARNING_FILL = "FFF7E6"


def set_run_font(run, size: int | None = None, bold: bool | None = None, color: RGBColor | None = None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None, size: int = 10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color or INK)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_widths(table, widths: list[float]):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    new_run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def style_document(doc: Document):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def add_title(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("Freight Intelligence")
    set_run_font(run, size=26, bold=True, color=RGBColor(11, 37, 69))

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("Microsoft Entra SSO 与服务器发布配置指南")
    set_run_font(run, size=15, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    run = p.add_run("适用系统：CourieDelivery / Freight Intelligence 运费预估、费率管理与账单复核平台")
    set_run_font(run, size=10, color=MUTED)


def add_callout(doc: Document, title: str, body: str, fill: str = CALLOUT_FILL):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_widths(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=11, bold=True, color=DARK_BLUE)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(body)
    set_run_font(r, size=10, color=INK)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_widths(table, widths)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_FILL)
        set_cell_text(cell, header, bold=True, color=DARK_BLUE, size=9)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=9)
    doc.add_paragraph()
    return table


def bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)


def number(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_title(doc)

    add_callout(
        doc,
        "核心结论",
        "要完成 Microsoft 一键登录或 silent login，建议先把系统发布到稳定服务器，并使用固定公网域名或固定外部访问地址。"
        "Microsoft Entra 的 Redirect URI 必须与浏览器实际访问地址完全一致，包括协议、域名和端口。生产环境建议使用 HTTPS 443；"
        "5173 和 8010 只作为开发端口，不建议对外开放。",
        WARNING_FILL,
    )

    doc.add_heading("1. 为什么需要先发布到服务器", level=1)
    for item in [
        "Entra Redirect URI 需要配置成用户浏览器实际访问系统的地址，例如 https://freight.company.com/。",
        "如果使用非标准端口，Redirect URI 也必须带端口，例如 https://freight.company.com:8443/。",
        "Silent login 依赖 Microsoft Entra 浏览器会话 cookie 和 MSAL 本地缓存；稳定域名有利于缓存和 SSO 行为保持一致。",
        "生产环境应使用 HTTPS；localhost 可用于开发，但不能代表正式 SSO 配置。",
    ]:
        bullet(doc, item)

    doc.add_heading("2. 推荐部署拓扑", level=1)
    add_table(
        doc,
        ["层级", "推荐做法", "说明"],
        [
            ["公网入口", "HTTPS 443 + DNS 域名", "推荐使用 https://freight.company.com/；80 仅用于跳转到 443。"],
            ["反向代理", "Nginx / IIS / Caddy", "统一代理前端静态文件和 /api 后端请求，隐藏内部端口。"],
            ["前端", "Vite build 后静态文件", "生产环境不运行 5173；只发布 frontend/dist。"],
            ["后端", "Django API 内网监听", "可监听 127.0.0.1:8010 或内网端口，由反向代理转发。"],
            ["数据库", "PostgreSQL 独立实例优先", "订单、SKU、报价、审计和账单数据会持续增长，建议与应用层分离。"],
        ],
        [1.2, 1.9, 3.4],
    )

    add_callout(
        doc,
        "推荐 URL 形态",
        "对外只暴露一个入口：https://freight.company.com/。前端访问根路径，API 通过 https://freight.company.com/api/ 转发。"
        "这样 Entra Redirect URI、CORS、Cookie/SSO 行为都会更稳定。",
    )

    doc.add_heading("3. 端口与网络开放建议", level=1)
    add_table(
        doc,
        ["端口", "用途", "是否对外开放", "建议"],
        [
            ["443", "HTTPS Web 入口", "是", "生产唯一推荐对外端口。"],
            ["80", "HTTP 跳转 HTTPS", "可选", "只做 301/308 跳转到 443，不承载登录。"],
            ["5173", "Vite 开发前端", "否", "仅本地开发使用，不用于生产。"],
            ["8010", "Django API 开发端口", "否", "生产只允许反向代理或内网访问。"],
            ["5432", "PostgreSQL", "否", "绝不公网开放，只允许应用服务器或 VPN 内网访问。"],
            ["1433", "SQL Server 数据源", "否", "只允许同步任务服务器访问 192.168.72.8。"],
        ],
        [0.8, 1.8, 1.2, 2.7],
    )

    doc.add_heading("4. Microsoft Entra SSO 配置步骤", level=1)
    for step in [
        "在 Microsoft Entra admin center 创建 App Registration，账号类型选择当前组织目录。",
        "Authentication 中添加 Single-page application 平台，并配置生产 Redirect URI，例如 https://freight.company.com/。",
        "Expose an API 中设置 Application ID URI，建议 api://<Application Client ID>。",
        "新增 scope：access_as_user，最终 scope 形如 api://<Application Client ID>/access_as_user。",
        "在前端 .env 中设置 VITE_MSAL_CLIENT_ID、VITE_MSAL_TENANT_ID、VITE_MSAL_SCOPE。",
        "在后端 .env 中设置 MSAL_TENANT_ID、MSAL_AUDIENCE，并在正式环境关闭 AUTH_ALLOW_DEV_USER。",
        "在系统 Admin -> Users & Roles 中把用户的 Entra Object ID / UPN / Tenant ID 与本地权限角色关联。",
    ]:
        number(doc, step)

    doc.add_heading("5. 当前系统需要的环境变量", level=1)
    add_table(
        doc,
        ["位置", "变量", "示例", "说明"],
        [
            ["frontend/.env", "VITE_API_BASE_URL", "https://freight.company.com/api", "生产前端调用后端 API 的地址。"],
            ["frontend/.env", "VITE_MSAL_CLIENT_ID", "<Application client ID>", "Entra App Registration 的客户端 ID。"],
            ["frontend/.env", "VITE_MSAL_TENANT_ID", "<Directory tenant ID>", "组织租户 ID。"],
            ["frontend/.env", "VITE_MSAL_SCOPE", "api://<client-id>/access_as_user", "前端请求本系统 API token 的 scope。"],
            ["backend/.env", "MSAL_TENANT_ID", "<Directory tenant ID>", "后端校验 token issuer。"],
            ["backend/.env", "MSAL_AUDIENCE", "api://<client-id>", "后端校验 token aud，必须匹配 API audience。"],
            ["backend/.env", "AUTH_ALLOW_DEV_USER", "False", "正式环境必须关闭开发管理员 fallback。"],
            ["backend/.env", "DJANGO_ALLOWED_HOSTS", "freight.company.com", "只允许正式域名访问。"],
            ["backend/.env", "CORS_ALLOWED_ORIGINS", "https://freight.company.com", "如果前后端同域，可保持最小化配置。"],
        ],
        [1.0, 1.55, 2.0, 1.95],
    )

    doc.add_heading("6. Silent login 的实际前提", level=1)
    for item in [
        "用户已经在当前浏览器中登录过 Microsoft 365 / Entra，或此前使用本系统完成过 Microsoft 登录。",
        "浏览器没有阻止 silent iframe 所需的第三方 cookie；Safari 或严格隐私设置可能需要交互登录。",
        "同一浏览器会话中不要同时存在多个 Microsoft 账号，否则 Entra 可能要求用户选择账号。",
        "加入 Entra 的 Windows 电脑会提高无感体验的概率，但 Web 应用仍主要依赖浏览器中的 Entra session cookie 和 MSAL cache。",
    ]:
        bullet(doc, item)

    doc.add_heading("7. 服务器规格建议", level=1)
    add_table(
        doc,
        ["方案", "适用场景", "应用服务器", "数据库服务器", "备注"],
        [
            [
                "测试 / UAT",
                "功能验证、小规模订单",
                "4 vCPU / 16 GB RAM / 100 GB SSD",
                "4 vCPU / 16-32 GB RAM / 300 GB SSD",
                "可与现有 PostgreSQL 共用，但不建议作为长期生产。",
            ],
            [
                "生产推荐",
                "百万行订单 + 日常同步 + 审计查询",
                "4-8 vCPU / 16-32 GB RAM / 150 GB SSD",
                "8 vCPU / 32-64 GB RAM / 1 TB NVMe SSD",
                "推荐应用与数据库分离，数据库使用高 IOPS 存储。",
            ],
            [
                "单机生产起步",
                "预算有限但需上线",
                "12 vCPU / 64 GB RAM / 1 TB NVMe SSD",
                "同机 PostgreSQL",
                "可先上线，后续数据增长后拆分数据库。",
            ],
            [
                "高增长 / 多年历史",
                "千万级明细、账单、报价轨迹",
                "8 vCPU / 32 GB RAM",
                "16 vCPU / 64-128 GB RAM / 2 TB+ NVMe",
                "建议读写分离、分区表、备份库和监控告警。",
            ],
        ],
        [0.9, 1.15, 1.55, 1.55, 1.35],
    )

    doc.add_heading("8. 百万行级别订单数据的数据库建议", level=1)
    for item in [
        "百万行订单本身不是 PostgreSQL 的上限，真正需要关注的是订单明细、tracking、quote_result、quote_trace_log、invoice_detail 等关联表会快速放大数据量。",
        "对历史订单、账单、报价审计结果建议按月份或财务期间分区，尤其是 quote trace 和 invoice detail。",
        "常用查询字段必须建立组合索引：order_no、platform_order_no、tracking_no、carrier_id、warehouse_id、postcode、created_at、invoice_date。",
        "使用增量同步而不是全量覆盖；同步任务记录 watermark，例如 external_updated_at 或 source_updated_at。",
        "报价和对账类批处理建议异步执行，避免在用户请求线程中直接处理数千或数万单。",
        "开启定期 VACUUM / ANALYZE，监控慢查询，并为大型导入设置批量写入和事务批次。",
        "备份容量至少按数据库有效数据的 2 倍预留；建议每日全量或增量备份，并定期做恢复演练。",
    ]:
        bullet(doc, item)

    doc.add_heading("9. PostgreSQL 参数起步建议", level=1)
    add_table(
        doc,
        ["参数", "32 GB RAM 数据库", "64 GB RAM 数据库", "说明"],
        [
            ["shared_buffers", "8 GB", "16 GB", "通常取内存约 25%。"],
            ["effective_cache_size", "24 GB", "48 GB", "帮助优化器估算 OS cache。"],
            ["work_mem", "16-32 MB", "32-64 MB", "按并发调整，不能盲目过大。"],
            ["maintenance_work_mem", "1 GB", "2 GB", "用于索引创建、VACUUM 等维护任务。"],
            ["max_connections", "100 或更低", "100 或更低", "建议配合 PgBouncer 做连接池。"],
            ["wal_compression", "on", "on", "降低 WAL 体积，适合大量导入更新。"],
        ],
        [1.7, 1.4, 1.4, 2.0],
    )

    doc.add_heading("10. 上线实施清单", level=1)
    for item in [
        "准备服务器、公网域名、HTTPS 证书和防火墙规则。",
        "确认生产访问地址，例如 https://freight.company.com/，并写入 Entra Redirect URI。",
        "构建前端：npm run build，将 frontend/dist 发布到 Web 根目录。",
        "部署后端：安装依赖、设置 .env、运行 python manage.py migrate。",
        "配置反向代理：/ 指向前端静态资源，/api 指向 Django API。",
        "关闭 AUTH_ALLOW_DEV_USER，设置强 DJANGO_SECRET_KEY，并限制 ALLOWED_HOSTS 和 CORS。",
        "在 Users & Roles 中建立管理员和关键用户的 Entra 关联。",
        "用公司账号测试 Microsoft 一键登录；再关闭浏览器重开测试 silent login。",
        "导入一批订单和账单数据，验证查询、报价、audit matrix 和 reconciliation 性能。",
        "配置数据库备份、日志轮转、磁盘空间、CPU/RAM/慢查询监控。",
    ]:
        number(doc, item)

    doc.add_heading("11. 安全与运维注意事项", level=1)
    for item in [
        "不要把 PostgreSQL 5432 或 SQL Server 1433 暴露到公网。",
        "生产环境所有登录和 API 流量必须走 HTTPS。",
        "不要在前端或文档中保存数据库密码；生产密码应使用服务器密钥管理或受控环境变量。",
        "后端必须校验 Entra token 的 issuer、signature 和 audience；当前系统通过 MSAL_TENANT_ID 与 MSAL_AUDIENCE 执行校验。",
        "保留本地管理员账号作为应急入口，但必须使用强密码并限制权限人员。",
        "对 invoice、order、quote audit 等导出功能保留审计日志。",
    ]:
        bullet(doc, item)

    doc.add_heading("12. 参考资料", level=1)
    refs = [
        (
            "Microsoft SPA app configuration",
            "https://learn.microsoft.com/en-us/entra/identity-platform/scenario-spa-app-configuration",
        ),
        ("MSAL.js single sign-on", "https://learn.microsoft.com/en-us/entra/identity-platform/msal-js-sso"),
        (
            "Microsoft identity platform access token validation",
            "https://learn.microsoft.com/en-us/entra/identity-platform/access-tokens",
        ),
    ]
    for title, url in refs:
        p = doc.add_paragraph(style="List Bullet")
        add_hyperlink(p, title, url)
        r = p.add_run(f" - {url}")
        set_run_font(r, size=10, color=MUTED)

    doc.core_properties.title = "Freight Intelligence SSO 与服务器发布配置指南"
    doc.core_properties.subject = "Microsoft Entra SSO, silent login, server deployment, PostgreSQL sizing"
    doc.core_properties.author = "Codex"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
