# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Freight_Intelligence_Intranet_PostgreSQL_Optimization_Guide.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(100, 116, 139)
HEADER_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
WARNING_FILL = "FFF7E6"


def set_run_font(run, size: int | None = None, bold: bool | None = None, color: RGBColor | None = None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None, size: int = 9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color or INK)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_widths(table, widths: list[float]):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    new_run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def style_document(doc: Document):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def paragraph(doc: Document, text: str, size: int = 11, color: RGBColor | None = None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size=size, color=color or INK)
    return p


def bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)


def number(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)


def add_callout(doc: Document, title: str, body: str, fill: str = CALLOUT_FILL):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_widths(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=11, bold=True, color=DARK_BLUE)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(body)
    set_run_font(r, size=10, color=INK)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_widths(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, HEADER_FILL)
        set_cell_text(cell, header, bold=True, color=DARK_BLUE, size=9)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=9)
    doc.add_paragraph()
    return table


def add_title(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Freight Intelligence")
    set_run_font(r, size=26, bold=True, color=RGBColor(11, 37, 69))

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("内网发布与 PostgreSQL 优化方案")
    set_run_font(r, size=15, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    r = p.add_run("前提：系统发布在公司内网 192.168.72.xx；PostgreSQL 数据库已完成；本文只提供发布与优化建议。")
    set_run_font(r, size=10, color=MUTED)


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_title(doc)

    add_callout(
        doc,
        "建议结论",
        "公司内网部署可以接 Microsoft Entra SSO。推荐使用内网 DNS 域名 + HTTPS，例如 https://freight-intelligence.company.local/，"
        "而不是直接暴露 http://192.168.72.xx。生产访问只开放 443；前端和后端通过反向代理统一成同一个站点；"
        "PostgreSQL 保持内网访问，不对公网开放。",
        WARNING_FILL,
    )

    doc.add_heading("1. 内网发布架构建议", level=1)
    add_table(
        doc,
        ["组件", "推荐配置", "优化原因"],
        [
            ["访问地址", "https://freight-intelligence.company.local/", "域名比 IP 更利于证书、SSO Redirect URI、缓存和后续迁移。"],
            ["对外端口", "443 HTTPS；80 可选跳转", "Microsoft Entra Redirect URI 推荐 HTTPS；生产不建议使用 HTTP。"],
            ["反向代理", "Nginx / IIS / Caddy", "统一前端静态文件和 /api，隐藏 5173/8010 等开发端口。"],
            ["前端发布", "frontend/dist 静态文件", "生产不运行 Vite dev server。"],
            ["后端 API", "Django 内网端口，例如 127.0.0.1:8010", "只允许反向代理访问，减少攻击面。"],
            ["PostgreSQL", "沿用已完成数据库，内网 5432", "只允许应用服务器和管理机访问。"],
        ],
        [1.2, 2.0, 3.3],
    )

    doc.add_heading("2. Microsoft Entra SSO 内网配置要点", level=1)
    for item in [
        "Microsoft 不需要从公网访问你的内网服务器；登录完成后是用户浏览器跳回你的内网 URL。",
        "Redirect URI 必须与浏览器实际访问地址完全一致，包括协议、域名/IP、端口和末尾路径。",
        "推荐 Redirect URI：https://freight-intelligence.company.local/；如果使用 IP，则必须是 https://192.168.72.xx/。",
        "如果使用 IP 证书，证书 Subject Alternative Name 必须包含该 IP；否则浏览器会提示证书不可信，SSO 体验会受影响。",
        "前端 VITE_MSAL_SCOPE 继续使用 api://<client-id>/access_as_user；后端 MSAL_AUDIENCE 使用 api://<client-id>。",
        "正式环境将 AUTH_ALLOW_DEV_USER=False，避免未带 token 时进入开发管理员 fallback。",
    ]:
        bullet(doc, item)

    doc.add_heading("3. 应用服务器推荐规格", level=1)
    add_table(
        doc,
        ["场景", "CPU / RAM", "磁盘", "适用建议"],
        [
            ["UAT / 展示", "4 vCPU / 16 GB", "100-200 GB SSD", "演示、SSO 验证、少量订单同步。"],
            ["生产起步", "8 vCPU / 32 GB", "200-500 GB SSD", "适合前端、Django API、同步任务和普通报价计算。"],
            ["批量审计较多", "8-12 vCPU / 32-64 GB", "500 GB SSD", "建议拆出 Celery/worker 或独立批处理进程。"],
            ["数据库同机临时方案", "12 vCPU / 64 GB", "1 TB NVMe", "可以起步，但长期建议应用和数据库分离。"],
        ],
        [1.3, 1.5, 1.35, 3.0],
    )

    doc.add_heading("4. PostgreSQL 已完成后的优先优化清单", level=1)
    add_table(
        doc,
        ["优先级", "优化项", "目标"],
        [
            ["P0", "慢查询监控 + pg_stat_statements", "先知道慢在哪里，再做索引和分区。"],
            ["P0", "关键查询组合索引", "支撑订单、tracking、账单、报价和 audit matrix 查询。"],
            ["P0", "连接池 PgBouncer / Django CONN_MAX_AGE", "避免大量短连接拖垮数据库。"],
            ["P1", "大表分区", "历史订单、invoice detail、quote trace、audit result 按月份或季度管理。"],
            ["P1", "批量同步 staging table", "ERP/WMS/InvoiceReader 数据先入 staging，再 merge/upsert。"],
            ["P1", "定期 VACUUM / ANALYZE", "保持查询计划准确，减少膨胀。"],
            ["P2", "物化视图 / 汇总表", "Dashboard、订单统计和对账汇总不直接扫明细表。"],
            ["P2", "冷热数据归档", "保留业务必要查询，压缩长期历史和 trace 日志成本。"],
        ],
        [0.8, 2.3, 3.4],
    )

    doc.add_heading("5. 建议索引方向", level=1)
    add_table(
        doc,
        ["数据域", "建议索引字段", "用途"],
        [
            ["订单主表", "erp_order_no, platform_order_no, rd3_order_id, platform_id, warehouse_id, created_at", "订单列表、平台筛选、历史查询。"],
            ["订单发货/tracking", "tracking_no, carrier_id, order_id, shipped_at", "Invoice matching 和 audit tracking 排序。"],
            ["订单 SKU 明细", "order_id, sku, sku_type, category", "批量报价、SKU 维度分析。"],
            ["Invoice header/detail", "invoice_source_id, invoice_date, tracking_no, carrier/service, batch_id", "按账单批次、来源、tracking 匹配实际费用。"],
            ["报价结果", "quote_run_id, order_id, tracking_no, carrier_id, status, created_at", "报价历史、audit matrix、异常排查。"],
            ["报价轨迹日志", "quote_result_id, order_id, tracking_no, created_at", "trace 页面和历史诊断。"],
            ["费率表", "carrier_id, service_id, warehouse_id, zone_code, effective_from, effective_to, is_active", "快速匹配当前有效 rate card/rule。"],
            ["SKU Master", "sku, sku_type, category", "Manual Quote 搜索和 SKU/Combo SKU 查询。"],
        ],
        [1.25, 3.25, 2.0],
    )
    add_callout(
        doc,
        "搜索优化补充",
        "列表页如果需要全库模糊搜索 SKU、订单号、tracking、平台订单号，建议启用 pg_trgm 扩展，并为高频文本字段创建 GIN trigram 索引。"
        "不要只在前端搜索当前页数据。",
    )

    doc.add_heading("6. 分区与归档建议", level=1)
    for item in [
        "historical_order、invoice_reconciliation_item、quote_trace_log、freight_audit_result 这类持续增长表，建议按 created_at / invoice_date 做月度或季度分区。",
        "quote_trace_log 体积增长最快，建议保留最近 6-12 个月明细；更早数据转 archive 表或压缩备份。",
        "invoice detail 和 audit result 保留业务追溯周期，例如 24-36 个月；超过周期后转只读归档。",
        "分区表需要配套本地索引，避免全分区扫描。",
        "历史回算时必须根据订单日期匹配对应 effective_from/effective_to 的 rate card，不要只取当前 active rate。",
    ]:
        bullet(doc, item)

    doc.add_heading("7. 批量同步与报价计算优化", level=1)
    for item in [
        "ERP/WMS/InvoiceReader 同步使用 watermark 增量字段，例如 updated_at、source_updated_at 或 invoice_date。",
        "批量导入建议每批 1,000-5,000 行，使用 bulk insert/upsert，避免逐行 save。",
        "大批量订单审计不要在 Web 请求中同步计算，应进入后台任务队列。",
        "Freight Audit Matrix 建议按 order 或 consignment 聚合后计算，避免同一订单多个 tracking 重复拉取 SKU 和地址数据。",
        "系统报价结果要保存 snapshot：SKU 尺寸、数量、地址、warehouse、platform、rate card version、breakdown 和 trace。",
        "API 型快递报价需要缓存 request/response 摘要，避免同一个订单重复调用外部 API。",
    ]:
        bullet(doc, item)

    doc.add_heading("8. PostgreSQL 参数起步建议", level=1)
    add_table(
        doc,
        ["参数", "32 GB RAM 数据库", "64 GB RAM 数据库", "说明"],
        [
            ["shared_buffers", "8 GB", "16 GB", "一般取内存约 25%。"],
            ["effective_cache_size", "24 GB", "48 GB", "帮助优化器估算可用缓存。"],
            ["work_mem", "16-32 MB", "32-64 MB", "按并发谨慎调大，避免内存爆掉。"],
            ["maintenance_work_mem", "1 GB", "2 GB", "用于建索引、VACUUM 等维护。"],
            ["max_connections", "100 或更低", "100 或更低", "配合 PgBouncer，避免连接过多。"],
            ["wal_compression", "on", "on", "降低大量导入/更新时 WAL 体积。"],
            ["autovacuum", "on，并调优阈值", "on，并调优阈值", "大表必须依赖 autovacuum 保持健康。"],
        ],
        [1.7, 1.4, 1.4, 2.0],
    )

    doc.add_heading("9. 备份、监控与维护", level=1)
    for item in [
        "开启每日备份，生产建议支持 PITR（Point-in-Time Recovery）。",
        "备份保留至少覆盖一个完整账单争议周期；建议不少于 30-90 天。",
        "监控 CPU、RAM、磁盘容量、IOPS、连接数、慢查询、表膨胀、autovacuum 状态。",
        "上线后每周审查 pg_stat_statements 前 20 条慢查询，并根据真实查询模式补索引。",
        "建立数据保留策略：quote trace、audit result、invoice detail 不应无限增长。",
        "对大索引维护使用 CREATE INDEX CONCURRENTLY / REINDEX CONCURRENTLY，避免长时间锁表。",
    ]:
        bullet(doc, item)

    doc.add_heading("10. 内网安全边界", level=1)
    add_table(
        doc,
        ["项目", "建议"],
        [
            ["数据库访问", "PostgreSQL 5432 仅允许应用服务器、DBA 管理机或 VPN 管理段访问。"],
            ["应用访问", "用户只访问 HTTPS 443；不要开放 Django 8010 或 Vite 5173。"],
            ["证书", "使用公司 CA 或受信任证书；如果用 IP，证书 SAN 必须包含 IP。"],
            ["账号", "关闭 AUTH_ALLOW_DEV_USER；保留强密码本地 break-glass 管理员。"],
            ["权限", "通过 Users & Roles 分配角色；导出、对账、费率管理建议保留审计日志。"],
            ["日志", "不要在日志中输出数据库密码、完整 access token 或敏感 invoice 原文。"],
        ],
        [1.55, 4.95],
    )

    doc.add_heading("11. 推荐实施顺序", level=1)
    for item in [
        "第一阶段：内网域名、HTTPS、反向代理、生产 .env、关闭 dev fallback。",
        "第二阶段：配置 Entra Redirect URI、API Scope、用户 Entra 关联，验证一键登录和 silent login。",
        "第三阶段：开启 pg_stat_statements，记录真实慢查询。",
        "第四阶段：补齐订单、tracking、invoice、quote、rate card 的关键组合索引。",
        "第五阶段：把历史订单、invoice detail、quote trace、audit result 规划成分区或 archive。",
        "第六阶段：把大批量同步和 audit matrix 计算迁移到后台任务队列。",
        "第七阶段：建立备份、恢复演练、慢查询巡检和容量告警。",
    ]:
        number(doc, item)

    doc.add_heading("12. 参考资料", level=1)
    refs = [
        (
            "Microsoft Redirect URI best practices and limitations",
            "https://learn.microsoft.com/en-us/entra/identity-platform/reply-url",
        ),
        ("MSAL.js single sign-on", "https://learn.microsoft.com/en-us/entra/identity-platform/msal-js-sso"),
        (
            "Microsoft access token validation",
            "https://learn.microsoft.com/en-us/entra/identity-platform/access-tokens",
        ),
        ("PostgreSQL pg_stat_statements", "https://www.postgresql.org/docs/current/pgstatstatements.html"),
        ("PostgreSQL table partitioning", "https://www.postgresql.org/docs/current/ddl-partitioning.html"),
    ]
    for title, url in refs:
        p = doc.add_paragraph(style="List Bullet")
        add_hyperlink(p, title, url)
        r = p.add_run(f" - {url}")
        set_run_font(r, size=10, color=MUTED)

    doc.core_properties.title = "Freight Intelligence 内网发布与 PostgreSQL 优化方案"
    doc.core_properties.subject = "Intranet deployment, Microsoft Entra SSO, PostgreSQL optimization"
    doc.core_properties.author = "Codex"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
